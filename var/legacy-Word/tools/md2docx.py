"""md -> docx 政府公文格式转换脚本

用法:
  python md2docx_python.py input.md [output.docx]

映射规则:
  #  标题        → 宋体 二号(22pt) 加黑 居中
  ## 一级标题    → 黑体 三号(16pt) 不加黑 首行缩进2字符
  ### 二级标题   → 楷体 三号(16pt) 加黑 首行缩进2字符
  #### 三级标题  → 仿宋 三号(16pt) 加黑 首行缩进2字符
  **前缀。**正文 / **前缀：**正文 / **前缀，**正文  → 段内混排（** 内加黑）
  **整段** / 普通**局部**继续                         → 同样按 ** 交替加黑
  <u>下划线</u>                                       → 下划线
  <span style="color:red">红字</span>                 → 字体红色
  <span style="background:#000;color:#fff">黑底</span> → 黑底白字
  *（注释）*     → 仿宋 无缩进
  <div align="right"> → 右对齐
  ---           → 禁止使用，跳过

页面设置: A4, 上下2.54cm, 左右3.17cm, 固定值30磅行距, 禁止孤行控制
页码: 宋体五号(10.5pt) 居中
"""
import sys
import os
import re
import html
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_indent(paragraph, chars):
    """chars=0:无缩进  chars=2:首行缩进2字符（字符单位，随字号自动缩放）

    直接设置 w:firstLineChars 属性（单位0.01字符，200=2字符），
    而不是 w:firstLine（固定磅值），这样修改字号时缩进自动跟随。"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = pPr.makeelement(qn('w:ind'), {})
        pPr.insert(0, ind)
    # 清除旧缩进，用字符单位替代磅值
    ind.attrib.pop(qn('w:firstLineChars'), None)
    ind.attrib.pop(qn('w:firstLine'), None)
    if chars > 0:
        ind.set(qn('w:firstLineChars'), str(chars * 100))
    else:
        ind.set(qn('w:firstLineChars'), '0')


def set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def apply_run_marks(run, marks):
    """下划线 / 红字 / 黑底白字等行内样式。"""
    if marks.get('underline'):
        run.underline = True
    color = marks.get('color')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if marks.get('black_bg'):
        run.font.highlight_color = WD_COLOR_INDEX.BLACK
        if not color:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _parse_span_style(style):
    """解析 span style → marks。"""
    marks = {}
    if not style:
        return marks
    compact = re.sub(r'\s+', '', style.lower())
    # 黑底（可带白字）
    if re.search(r'background(?:-color)?:(?:#000(?:000)?|black)\b', compact):
        marks['black_bg'] = True
        if re.search(r'color:(?:#fff(?:fff)?|white)\b', compact):
            marks['color'] = 'FFFFFF'
        return marks
    # 红字
    if re.search(r'color:(?:red|#f00(?:0)?|#ff0000)\b', compact):
        marks['color'] = 'FF0000'
    return marks


def iter_inline_segments(text):
    """将正文拆成 (纯文本, marks) 段。支持 **加黑** / <u> / <span style=…>。"""
    patterns = [
        ('bold', re.compile(r'\*\*(.+?)\*\*', re.S)),
        ('strong', re.compile(r'<(strong|b)>(.*?)</\1>', re.I | re.S)),
        ('u', re.compile(r'<u>(.*?)</u>', re.I | re.S)),
        ('span', re.compile(r'<span\s+([^>]*)>(.*?)</span>', re.I | re.S)),
    ]
    pos = 0
    n = len(text)
    while pos < n:
        best = None
        for kind, rx in patterns:
            m = rx.search(text, pos)
            if m and (best is None or m.start() < best[0]):
                best = (m.start(), m.end(), kind, m)
        if best is None:
            if pos < n:
                yield text[pos:], {}
            break
        start, end, kind, m = best
        if start > pos:
            yield text[pos:start], {}
        if kind == 'bold':
            yield m.group(1), {'bold': True}
        elif kind == 'strong':
            for sub, mk in iter_inline_segments(m.group(2)):
                yield sub, {**mk, 'bold': True}
        elif kind == 'u':
            for sub, mk in iter_inline_segments(m.group(1)):
                yield sub, {**mk, 'underline': True}
        elif kind == 'span':
            attrs, inner = m.group(1), m.group(2)
            style = ''
            sm = re.search(r'style\s*=\s*["\']([^"\']*)["\']', attrs, re.I)
            if sm:
                style = sm.group(1)
            base = _parse_span_style(style)
            for sub, mk in iter_inline_segments(inner):
                yield sub, {**base, **mk}
        pos = end


def add_inline_runs(paragraph, text, font_name, font_size):
    """正文行内：加黑 / 下划线 / 红字 / 黑底白字。"""
    text = _normalize_export_text(text) if text else ""
    if not text:
        return
    for seg, marks in iter_inline_segments(text):
        if not seg:
            continue
        # 段内残留实体再解一次（标题规则已解过整行时无害）
        seg = html.unescape(seg).replace("\u00a0", " ")
        if not seg:
            continue
        run = paragraph.add_run(seg)
        set_font(run, font_name, font_size, bool(marks.get('bold')))
        apply_run_marks(run, marks)


def set_line_spacing(paragraph, rule=3, value=30, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = rule  # 3 = 固定值
    pf.line_spacing = Pt(value)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def set_paragraph_style(paragraph, indent_chars=None):
    """设置段落样式：缩进 + 孤行控制 + 首位字符控制（中文习惯）"""
    pPr = paragraph._element.get_or_add_pPr()

    # 孤行控制 + 首位字符控制
    jc = pPr.find(qn('w:suppressLine'))
    if jc is None:
        jc = pPr.makeelement(qn('w:suppressLine'), {})
        pPr.insert(0, jc)
    jc.set(qn('w:val'), '0')  # 禁止行首字符孤行

    # 首行缩进（字符单位）
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = pPr.makeelement(qn('w:ind'), {})
        pPr.insert(0, ind)
    ind.attrib.pop(qn('w:firstLineChars'), None)
    ind.attrib.pop(qn('w:firstLine'), None)
    if indent_chars is not None and indent_chars > 0:
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))
    else:
        ind.set(qn('w:firstLineChars'), '0')


def add_para(doc, text, font_name, font_size, bold, alignment, indent_chars=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_name, font_size, bold)
    p.alignment = alignment
    set_line_spacing(p)
    set_paragraph_style(p, indent_chars)
    return p, run


def _set_cell_shading(cell, fill):
    """单元格底纹（表头浅灰）"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, size=12):
    """单元格文字：仿宋、居中、垂直居中"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '仿宋', size, bold)
    set_line_spacing(p)
    set_paragraph_style(p, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_row_no_break(row, min_height_pt=None):
    """禁止行跨页断开，可选设置最小行高"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    # 禁止行跨页
    cant_split = OxmlElement('w:cantSplit')
    trPr.append(cant_split)
    # 最小行高
    if min_height_pt is not None:
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(min_height_pt * 20)))  # 1pt = 20 twips
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)


def _set_row_header(row):
    """设为表头行（跨页重复）"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    trPr.append(tbl_header)


def _set_keep_next(paragraph):
    """段落与下一段保持同页（keepNext）"""
    pPr = paragraph._element.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)


def _set_cant_split_table(table):
    """禁止表格跨页断表"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    cant_split = OxmlElement('w:cantSplit')
    tblPr.append(cant_split)


def add_table(doc, rows, last_para=None, force_page_break=False):
    """rows: 二维列表，首行为表头 → Word 表格（Table Grid、表头加粗灰底）

    处理规则：
    - 禁止行跨页断开
    - 首行为表头跨页重复
    - 最小行高约0.8cm
    - 禁止表格整体跨页断表
    - 大表格（8行以上）前加分页符，确保整表从新页开始
    - 表头说明段落跟随表格（keepNext）

    last_para: 表格前一段落（表头说明），自动设 keepNext
    force_page_break: True 时在表格前插入分页符
    """
    nrows = len(rows)
    ncols = max(len(r) for r in rows)

    # 大表格前加分页符
    if force_page_break or nrows >= 8:
        doc.add_page_break()

    # 表头说明段落跟随表格
    if last_para is not None:
        _set_keep_next(last_para)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(nrows):
        row = table.rows[r]
        # 禁止行跨页 + 最小行高
        _set_row_no_break(row, min_height_pt=24)
        for c in range(ncols):
            text = rows[r][c] if c < len(rows[r]) else ''
            cell = row.cells[c]
            is_header = (r == 0)
            _set_cell_text(cell, text, bold=is_header, size=12)
            if is_header:
                _set_cell_shading(cell, 'D9D9D9')
                _set_row_header(row)
    # 表格后留一空行
    sp = doc.add_paragraph()
    set_line_spacing(sp)
    return table


def _setup_document():
    """新建公文页：A4、页边距、Normal=仿宋三号。"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    return doc, section


def _add_page_number(section):
    """页脚居中页码：宋体五号。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run()
    fldChar1 = run_footer._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_footer._r.append(fldChar1)
    instrText = run_footer._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    run_footer._r.append(instrText)
    fldChar2 = run_footer._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_footer._r.append(fldChar2)
    set_font(run_footer, '宋体', 10.5, False)


def _try_consume_table(doc, lines, i, t):
    """若当前行起是 GFM 表格则写入并返回新下标；否则返回 None。"""
    if not re.match(r'^\|.*\|\s*$', t):
        return None
    block = [t]
    i += 1
    while i < len(lines) and re.match(r'^\|.*\|\s*$', lines[i].rstrip('\n').strip()):
        block.append(lines[i].rstrip('\n').strip())
        i += 1
    is_table = (len(block) >= 2 and all(
        re.match(r'^[-:\s]*$', c) for c in block[1].strip('|').split('|')))
    if not is_table:
        return None
    parsed = [[c.strip() for c in b.strip('|').split('|')] for b in block]
    parsed = [row for row in parsed
              if not all(re.match(r'^[-:\s]*$', c) for c in row)]
    if parsed:
        last_para = doc.paragraphs[-1] if doc.paragraphs else None
        add_table(doc, parsed, last_para=last_para, force_page_break=len(parsed) >= 8)
    return i


def _normalize_export_text(t):
    """编辑器空段会落盘为字面量 &nbsp;；导出前解码，避免 Word 里出现英文实体。"""
    s = html.unescape(str(t or ''))
    s = (
        s.replace("\u00a0", " ")
        .replace("\u2003", " ")
        .replace("\u2002", " ")
        .replace("\ufeff", "")
        .replace("\u200b", "")
    )
    # 整行只是空段占位 → 当作空行跳过
    if not s.strip():
        return ""
    return s


def _emit_md_line(doc, t):
    """处理单行非表格 md，写入 doc。"""
    t = _normalize_export_text(t)
    if t == "" or t == "---":
        return
    rules = [
        (r'^# (.+)$', lambda m: (
            add_para(doc, _normalize_export_text(m.group(1)), '宋体', 22, True, WD_ALIGN_PARAGRAPH.CENTER, 0),
            set_line_spacing(doc.add_paragraph()),
        )),
        (r'^## (.+)$', lambda m: add_para(
            doc, _normalize_export_text(m.group(1)), '黑体', 16, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 2)),
        (r'^### (.+)$', lambda m: add_para(
            doc, _normalize_export_text(m.group(1)), '楷体', 16, True, WD_ALIGN_PARAGRAPH.JUSTIFY, 2)),
        (r'^#### (.+)$', lambda m: add_para(
            doc, _normalize_export_text(m.group(1)), '仿宋', 16, True, WD_ALIGN_PARAGRAPH.JUSTIFY, 2)),
        (r'^\*（(.+)）\*$', lambda m: add_para(
            doc, f'（{_normalize_export_text(m.group(1))}）', '仿宋', 16, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0)),
        (r'^<div align="right">(.+)</div>$', lambda m: add_para(
            doc, _normalize_export_text(m.group(1)), '仿宋', 16, False, WD_ALIGN_PARAGRAPH.RIGHT, 0)),
    ]
    for pat, action in rules:
        m = re.match(pat, t)
        if m:
            action(m)
            return
    p = doc.add_paragraph()
    add_inline_runs(p, t, '仿宋', 16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p)
    set_paragraph_style(p, 2)


def convert(md_path, docx_path=None):
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + '.docx'
    if not os.path.exists(md_path):
        print(f'ERROR: File not found: {md_path}')
        sys.exit(1)

    doc, section = _setup_document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        t = lines[i].rstrip('\n').strip()
        nxt = _try_consume_table(doc, lines, i, t)
        if nxt is not None:
            i = nxt
            continue
        i += 1
        _emit_md_line(doc, t)

    _add_page_number(section)
    doc.save(docx_path)
    print(f'DONE: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    md_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else None
    convert(md_file, docx_file)
